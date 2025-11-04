"""
Utility for extracting text from DOCX and PDF files
"""
import os
from typing import Optional
from docx import Document


class DocumentProcessor:
    """Handles text extraction from various document formats"""

    @staticmethod
    def extract_text_from_docx(file_path: str) -> str:
        """
        Extract text from a DOCX file

        Args:
            file_path: Path to the DOCX file

        Returns:
            Extracted text as a string
        """
        try:
            doc = Document(file_path)
            text_parts = []

            # Extract text from paragraphs
            for paragraph in doc.paragraphs:
                if paragraph.text.strip():
                    text_parts.append(paragraph.text.strip())

            # Extract text from tables
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        if cell.text.strip():
                            text_parts.append(cell.text.strip())

            return '\n'.join(text_parts)
        except Exception as e:
            raise Exception(f"Error extracting text from DOCX: {str(e)}")

    @staticmethod
    def extract_text_from_pdf(file_path: str) -> str:
        """
        Extract text from a PDF file

        Args:
            file_path: Path to the PDF file

        Returns:
            Extracted text as a string
        """
        raise NotImplementedError("PDF extraction not yet implemented")

    @staticmethod
    def extract_text(file_path: str) -> str:
        """
        Extract text from a document based on its extension

        Args:
            file_path: Path to the document

        Returns:
            Extracted text as a string
        """
        _, ext = os.path.splitext(file_path)
        ext = ext.lower()

        if ext == '.docx':
            return DocumentProcessor.extract_text_from_docx(file_path)
        elif ext == '.pdf':
            return DocumentProcessor.extract_text_from_pdf(file_path)
        else:
            raise ValueError(f"Unsupported file format: {ext}")

    @staticmethod
    def clean_text(text: str) -> str:
        """
        Clean and normalize extracted text

        Args:
            text: Raw text to clean

        Returns:
            Cleaned text
        """
        # Remove excessive whitespace
        lines = [line.strip() for line in text.split('\n') if line.strip()]
        text = '\n'.join(lines)

        # Remove multiple consecutive newlines
        while '\n\n\n' in text:
            text = text.replace('\n\n\n', '\n\n')

        return text.strip()
